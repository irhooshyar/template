import subprocess
import math
import os
from pathlib import Path

import docx2txt
import glob
from hazm import *

from abdal import config
import threading
from abdal.settings import LOCAL_SETTING
import PyPDF2 

from pdf2image import convert_from_path
import pytesseract
from PIL import Image


from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import tempfile
import requests
import textract

def readFiles_parallel(path, readContent=True, preprocess=True, preprocessArg={}):

    numbers_of_thread = LOCAL_SETTING['NUMBER_OF_THREAD']
    result_list = [None] * numbers_of_thread

    def splitArray(arr, n):
        return_li = []

        step = math.ceil(arr.__len__() / n)
        for i in range(n):
            start_idx = i * step
            end_idx = min(start_idx + step, arr.__len__())
            return_li.append(arr[start_idx:end_idx])

        return return_li

    def body(files, thread_number):
        result_text = {}
        t = 1
        for file in files:
            if readContent:
                # read
                if file[-3:] == 'txt':
                    text = open(file, encoding="utf8").read()
                else:
                    text = docx2txt.process(file)

                # Arabic char convert
                text = arabicCharConvert(text)
                if preprocess:
                    text = Preprocessing(text, **preprocessArg)
            else:
                text = ""
            file = str(os.path.basename(file)).split(".")[0]
            result_text[file] = text

    #     return
        result_list[thread_number] = result_text

    def parallel_run(func, li, arg=()):
        thread_obj = []
        thread_number = 0
        for S in li:
            thread = threading.Thread(target=func,
                                      args=(S, thread_number)+arg)
            thread_obj.append(thread)
            thread_number += 1
            thread.start()

        for thread in thread_obj:
            thread.join()


    all_files = glob.glob(path + "//*.docx") + glob.glob(path + "//*.txt")
    all_files = splitArray(all_files, numbers_of_thread)
    parallel_run(body, all_files)

    # join res
    result = {}
    for r in result_list:
        result = {**result, **r}

    return result

def readFiles(path, readContent=True, preprocess=True, preprocessArg={}):
    # all_files = glob.glob(path + "//*.docx")
    all_files = glob.glob(path + "/*")
    result_text = {}
    c = 0
    all_files_count = len(all_files)
    for file in all_files:
        format = str(os.path.basename(file)).split(".")[-1]
        try:
            if readContent:
                text = ''
                if format.lower() == "doc":
                    text = textract.process(file, encoding='utf-8').decode("utf-8")
                    c += 1
                    print(f'{c}/{all_files_count}')
                if format.lower() =='docx':
                    text = docx2txt.process(file)
                # Arabic char convert
                text = arabicCharConvert(text)
                if preprocess:
                    text = Preprocessing(text, **preprocessArg)
            else:
                text = ""
            file = str(os.path.basename(file)).split(".")[0]
            result_text[file] = text
        except:
            print(f'{file} can not read')
        
    all_files = glob.glob(path + "//*.txt")
    for file in all_files:
        if readContent:
            text = open(file, encoding="utf8").read()
            # Arabic char convert
            text = arabicCharConvert(text)
            if preprocess:
                text = Preprocessing(text, **preprocessArg)
        else:
            text = ""
        file = str(os.path.basename(file)).split(".")[0]
        result_text[file] = text

    return result_text

def read_raw_text(path):

    all_files = glob.glob(path + "//*.docx")
    result_text = {}
    for file in all_files:
        text = docx2txt.process(file)

        file = str(os.path.basename(file)).split(".")[0]
        result_text[file] = text

    all_files = glob.glob(path + "//*.txt")
    for file in all_files:
        text = open(file, encoding="utf8").read()

        file = str(os.path.basename(file)).split(".")[0]
        result_text[file] = text

    return result_text

def readFile(filePath, preprocess=True, preprocessArg={}):
    if str(filePath).split(".")[-1] == "docx":
        text = docx2txt.process(filePath)
        # Arabic char convert
        text = arabicCharConvert(text)
        if preprocess:
            text = Preprocessing(text, **preprocessArg)
    elif str(filePath).split(".")[-1] == "txt":
        text = open(filePath,encoding="utf8").read()
        # Arabic char convert
        text = arabicCharConvert(text)
        if preprocess:
            text = Preprocessing(text, **preprocessArg)

    return text

def pdf2txt(dataPath,file_old,format):
    address = dataPath + "/" + file_old + format
    new_address = dataPath + "/" + file_old + '.txt'
    text = ""   
    pdfFileObj = open(address, 'rb') 
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 
        
    for i in range(pdfReader.numPages): 
        pageObj = pdfReader.getPage(i) 
        text += pageObj.extractText()
    # new_text = ""
    # for word in text.split():
    #     new_text += word[::-1]+" "
    pdfFileObj.close() 
    os.remove(address) 
    new_file = open(new_address, 'w', encoding="utf8")
    new_file.write(text[::-1])

    return


def pdf2txt2(dataPath,file_old,format):
    address = dataPath + "/" + file_old + format
    new_address = dataPath + "/" + file_old + '.txt'
    pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
    
    images = convert_from_path(address,poppler_path=r"D:\\poppler-0.68.0\\bin")
    text = ""
    for pg, img in enumerate(images):
        text += pytesseract.image_to_string(img,lang='fas')
        print(text)
        
    os.remove(address) 
    new_file = open(new_address, 'w', encoding="utf8")
    new_file.write(text)

    return


def pdf2txt3(dataPath, file_old, format):
    """ 
    A way to use google-tesseract ocr for extracting 
    texts from pdf file.
    
    Args:
        dataPath (str): PDF file path.
        output_dir (str): Output directory.
        lang (str): tesseract language support.
    """
    pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
    lang='fas'
    address = dataPath + "/" + file_old + format
    new_address = dataPath + "/" + file_old + '.txt'
        
    pdf_name = file_old
    pages = convert_from_path(address, poppler_path=r"D:\\poppler-0.68.0\\bin")
    texts = []
    
    for i, page in tqdm(enumerate(pages), position=0):
        
        with tempfile.TemporaryDirectory() as img_dir:
            img_name = f'{pdf_name}-{i+1}.jpg'
            img_path = os.path.join(img_dir, img_name)
            
            page.save(img_path, 'JPEG')
            text = pytesseract.image_to_string(Image.open(img_path), lang=lang)
            texts.append(text)
    
    document = Document()
    style_normal = document.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.rtl = True
    
    style_h1 = document.styles['Heading 1']
    font = style_h1.font
    font.name = 'Arial'
    font.rtl = True
    
    for i, text in tqdm(enumerate(texts), position=0):
        heading = document.add_heading(f'صفحه: {i+1}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.style = document.styles['Heading 1']
        
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.style = document.styles['Normal']
    
    output_path = os.path.join(dataPath, f'{pdf_name}.txt')
    document.save(output_path)
        
        
def pdf2txt0(dataPath,file_old,format,folder_name):
    address = dataPath + "/" + file_old + format
    new_address = dataPath + "/" + file_old + '.txt'
    download_link = 'http://37.156.144.109:8000/' + 'download_book/' + folder_name + '/' + file_old + '.' + format + '/'
    url = "https://www.eboo.ir/api/ocr/getway"
    token = "UtSSGtfjCZZ4INNqIme6wUUNGRWpO5s0"
    command = "addfile"

    myobj = {'token': token , 'command': command , 'filelink': download_link}

    x = requests.post(url, json = myobj).json()

    filetoken = x['FileToken']
    command = "convert"
    method = 1
    
    myobj = {'token': token , 'filetoken': filetoken , 'command': command, 'method': method}

    x = requests.post(url, json = myobj).json()

    URL = x['FileToDownload']
    response = requests.get(URL)
    res = response.content
    res = res.decode('utf-8')
    os.remove(address) 
    new_file = open(new_address, 'w')
    new_file.write(res)
    
    return
  
def convert_all_pdfs_to_txt(folder_name):
    dataPath = str(Path(config.DATA_PATH, folder_name))
    all_files = glob.glob(dataPath + "/*.pdf")
    for file in all_files:
        file_name = str(os.path.basename(file))
        format = "." + str(os.path.basename(file)).split(".")[-1]
        file_old = str(os.path.basename(file))[:-len(format)]
        # pdf2txt(dataPath,file_old,format)
        pdf2txt0(dataPath,file_old,format,folder_name)

        
def renameFilesToStandard(folder_name):
    dataPath = str(Path(config.DATA_PATH, folder_name))
    # all_files = glob.glob(dataPath + "/*")
    # for file in all_files:
    #     format = "." + str(os.path.basename(file)).split(".")[-1]
    #     file_old = str(os.path.basename(file))[:-len(format)]
    #     if format.lower() in ['.doc', '.docx']:
    #         # subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file_old + format])
    #         os.rename(Path(dataPath, file_old + format), Path(dataPath, file_old + '.docx'))
    
    # all_files = glob.glob(dataPath + "/*.docx") + glob.glob(dataPath + "/*.txt")
    all_files = glob.glob(dataPath + "/*")
    for file in all_files:
        try:
            format = "." + str(os.path.basename(file)).split(".")[-1]
            if format.lower() not in ['.doc', '.docx', '.txt']:
                continue
            file_old = str(os.path.basename(file))[:-len(format)]
            file_new = standardFileName(file_old)
            if file_old != file_new:
                os.rename(Path(dataPath, file_old + format), Path(dataPath, file_new + format))
        except:
            continue


def getStemDict(path, stem=False, remove_sw=False):
    all_files = readFiles(path, preprocessArg={"stem": stem, "removeSW": remove_sw})
    stem_dict = {}
    for file in all_files:
        text = all_files[file]
        for word in text:
            word_s = stemming(word)
            stem_dict[word] = word_s
    return stem_dict


def standardIndexName(Country,model_name):

    file_name = Country.file_name

    index_name = file_name.split('.')[0] + '_' + model_name
    index_name = index_name.replace(' ','_')
    index_name = index_name.replace(',','_')
    index_name = index_name.lower()

    return index_name

def standardFileName(name):
    name = name.replace(".", "")
    name = arabicCharConvert(name)
    name = persianNumConvert(name)
    name = name.strip()

    while "  " in name:
        name = name.replace("  "," ")

    return name

def persianNumConvert(text):
    persian_num_dict = {"۱": "1" ,"۲": "2", "۳": "3", "۴": "4", "۵": "5", "۶": "6", "۷": "7", "۸": "8", "۹":"9" , "۰": "0" }
    for key, value in persian_num_dict.items():
        text = text.replace(key, value)
    return text


def arabicCharConvert(text):
    arabic_char_dict = {"ى": "ی" ,"ك": "ک", "آ": "ا", "أ": "ا", "إ": "ا", "ي": "ی", "ة": "ه", "ۀ": "ه", "  ":" ", "\n\n":"\n", "\n ":"\n" , }
    for key, value in arabic_char_dict.items():
        text = text.replace(key, value)

    return text


stemmer = Stemmer()


def stemming(word):
    word_s = stemmer.stem(word)
    return word_s


def Preprocessing(text, tokenize=True, stem=True, removeSW=True, normalize=True, removeSpecialChar=True):

    # Cleaning
    if removeSpecialChar:
        ignoreList = ["!", "@", "$", "%", "^", "&", "*", "(", ")", "_", "+", "-", "/", "*", "'", "،", "؛", ",", ""
                                                                                                                "{",
                      "}", '\xad', '­'
                      "[", "]", "«", "»", "<", ">", ".", "?", "؟", "\n", "\t", '"',
                      '۱', '۲', '۳', '۴', '۵', '۶', '۷', '۸', '۹', '۰', "٫",
                      '0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
        for item in ignoreList:
            text = text.replace(item, " ")

    # Normalization
    if normalize:
        normalizer = Normalizer()
        text = normalizer.normalize(text)

        # Delete non-ACII char
        for ch in text:
            if ord(ch) <= 255 or (ord(ch) > 2000):
                text = text.replace(ch, " ")

    # Tokenization
    if tokenize:
        text = [word for word in text.split(" ") if word != ""]

        # stopwords
        if removeSW:
            stopwords_list = open(Path(config.BASE_PATH, "text_files/stopWords.txt"), encoding="utf8").read().split(
                "\n")
            text = [word for word in text if word not in stopwords_list]

            # filtering
            text = [word for word in text if len(word) >= 2]

        # stemming
        if stem:
            text = [stemming(word) for word in text]

    return text
